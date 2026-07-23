import io

from docx import Document
from docx.shared import RGBColor

from app.db.connection import get_connection
from app.errors import ApiError
from app.services.chart_service import get_rows
from app.services.session_service import get_session

_TIER_COLORS = {
    "Strong": RGBColor(0x1A, 0x7F, 0x37),
    "Moderate": RGBColor(0xB0, 0x8A, 0x00),
    "Weak": RGBColor(0xB0, 0x2A, 0x2A),
}


def build_docx(sid: str) -> bytes:
    conn = get_connection()
    try:
        session = get_session(sid, conn)
        if not session["generated"]:
            raise ApiError(400, "not_generated_yet", "Run Generate before exporting.")
    finally:
        conn.close()

    rows = get_rows(sid)

    document = Document()
    document.add_heading("Claim Chart", level=1)

    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].text = "Claim Element"
    header[1].text = "Evidence"
    header[2].text = "Reasoning"
    header[3].text = "Confidence"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["claim_element"]
        cells[1].text = row["product_feature"]
        cells[2].text = row["ai_reasoning"]
        confidence = row["confidence"] or ""
        run = cells[3].paragraphs[0].add_run(confidence)
        if confidence in _TIER_COLORS:
            run.font.color.rgb = _TIER_COLORS[confidence]

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
